# coding: utf-8
'''
用于获取django项目中注册过的app以及每个app下的models，生成数据字典
this script be used get all of apps&models of django, and create a library of model
'''
import docx
import json
import os
import functools
from django.db.models.fields.related import ForeignKey, OneToOneField, ManyToManyField

def parseModels(models):
    body = None
    with open(models, 'r') as f:
        body = [x.replace('\n', '').split('\t') for x in f.readlines()]
    group = {}
    for m in body:
        _group = group.get(m[0], [])
        _group.append(m[1])
        group.update({m[0]: _group})
    return group


def getAll(path, settings, models):

    exec('from %s import INSTALLED_APPS' % settings)
    from django.db.models.base import ModelBase
    apps = tuple(filter(lambda x: not x.startswith('django'), locals().get('INSTALLED_APPS')))
    models = open(models, 'a', encoding='utf-8')
    for app in apps:
        print('查询到已注册APP：', app)
        models_dir_dpath = os.path.join(os.path.abspath('.'), path, app, 'models')
        models_file_dpath = os.path.join(os.path.abspath('.'), path, app, 'models.py')
        model_path = []
        if os.path.isdir(models_dir_dpath):
            files = os.listdir(models_dir_dpath)
            for file in files:
                if file.startswith('_'):
                    continue
                if file.startswith('#'):
                    continue
                if file.endswith('pyc'):
                    continue
                model_path.append('.'.join([app, 'models', file]).replace('.py', ''))
        if os.path.isfile(models_file_dpath):
            model_path.append('.'.join([app, 'models']))
        print('APP所包含的model：', model_path)
        for model in model_path:
            exec('from %s import *' % model)
            for k, v in locals().items():
                if str(v).find(model)>0 and isinstance(v, ModelBase):
                    _t = str(v).split("'")[1]
                    _pack = _t.rsplit('.', 1)[0]
                    _cls = _t.rsplit('.', 1)[1]
                    models.write("\t".join([_pack, _cls]))
                    models.write("\n")
    models.close()




def getInfo(model):
    if not hasattr(model, '_meta'):
        return {'class_name': model.__name__}
    tb_db_name = model._meta.db_table
    tb_verbose_name = model._meta.verbose_name
    fields = model._meta.fields
    rel = list()
    for i in fields:
        # column: (type, verbose_name)
        column = i.column
        verbose_name = i.verbose_name
        _type = type(i)
        relevance = '无'
        if _type in (ForeignKey, OneToOneField, ManyToManyField):
            relevance = str(i.related_model).split("'")[1]
        rel.append({'field': column, 'alias': verbose_name, 'type': _type.__name__, 'relevance': relevance})
        # rel.append((column, verbose_name, _type.__name__, releated))
    _model = {'class_name': model.__name__, 'db_name': tb_db_name, 'verbose_name': tb_verbose_name, 'rel': rel}
    return _model

# exec 无返回值
# eval 可以获取返回值
def saveRel(group, table):
    with open(table, 'w', encoding='utf-8') as f:
        for k, v in group.items():
            exec('from %s import %s' % (k, ','.join(v)))
            fields = {}
            for _v in v:
                model = fields.get(k, [])
                model.append(eval('getInfo(%s)' % _v))
                fields[k] = model
            # print(fields)
            fields = json.dump(fields, f, ensure_ascii=False)
            f.write('\n')

# control docx
def createDoc(tables, _fileName):
    doc = docx.Document()
    for table in tables:
        table = json.loads(table)
        for model, fields in table.items():  # 遍历所有的应用
            doc.add_heading(model.title(), 1)  # 创建一级标题
            for field in fields:  # 遍历每个应用下的model
                _class = field.get('class_name')
                _db_name = field.get('db_name', False)
                _verbose_name = field.get('verbose_name')
                _rels = field.get('rel')
                doc.add_heading(_class, 2)  # 设置model名称
                if not _db_name:
                    continue
                p = '数据库：' + _db_name + '，别名：' + _verbose_name
                doc.add_paragraph(p)
                title = ['Column', 'Type', 'Relevance', 'Alias']  # 在doc中显示的标题
                relKeys = ['field', 'type', 'relevance', 'alias']  # 与之对应的字段，有关联，修改多处
                table = doc.add_table(rows=1, cols=4, style="Table Grid")  # 为每个model创建一个表格
                th = table.rows[0].cells  # 设置标题
                for i in range(4):
                    th[i].text = title[i]
                for rel in _rels:
                    td = table.add_row().cells
                    for i in range(4):
                        td[i].text = rel.get(relKeys[i])
    project_name = (os.path.abspath('.').rsplit(os.path.sep, 1)[1] + '.docx') if not _fileName else _fileName
    doc.save(project_name)

def getAllTables():
    tables = None
    with open('table', 'r', encoding='utf-8') as f:
        tables = f.readlines()
    return tables


'''
程序入口:
path: 包含所有app的文件夹目录，默认在根目录下，可为空
file: 生成docx的文件名，需要包含后缀，默认生成的为该项目名
settings: 注册app的settings文件位置，默认在根目录下，如有变动，则可以用相对路径，用点分割
models: 存放所有app以及其中的model
table: 存放表结构，数据内容为json

程序默认会生成2个文件，一个是收集的所有app以及model的关系，一个是基于前一个文件获取的表结构
默认情况下，文件名分别为：models.txt 以及 table
'''
def run(path='', settings='settings', file=None, models='models', table='table'):
    getAll(path=path, settings=settings, models=models)
    group = parseModels(models)
    saveRel(group, table)
    tables = getAllTables()
    createDoc(tables, file)
    # 清理工作，缓存文件删除
    os.remove(models)
    os.remove(table)
